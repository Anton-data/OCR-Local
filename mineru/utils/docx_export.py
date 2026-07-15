# Copyright (c) Opendatalab. All rights reserved.
"""Export a MinerU ``content_list.json`` recognition result into a Word (.docx) document.

This module is intentionally dependency-light (``python-docx`` + stdlib ``html.parser``)
so it can be reused as-is both from the Gradio UI and from the FastAPI service, and it is
also safe to import from external projects that embed MinerU as a library.

Supported ``content_list`` block types:

- ``text``: plain paragraph, or a heading when ``text_level`` (1..N) is present.
- ``table``: ``table_body`` is an HTML table; it is parsed with the stdlib
  ``html.parser`` and rendered as a native Word table (style ``Table Grid``),
  with ``rowspan``/``colspan`` handled via cell merging. ``table_caption`` /
  ``table_footnote`` are rendered as paragraphs before/after the table.
- ``image``: the image referenced by ``img_path`` is embedded. ``image_caption`` /
  ``image_footnote`` are rendered as paragraphs before/after the picture.
- ``equation``: rendered as plain text (the raw LaTeX source); no LaTeX rendering
  is performed.

A handful of additional block types commonly emitted by MinerU (``header``,
``footer``, ``page_number``, ``aside_text``, ``ref_text``, ``list``) are also
rendered as best-effort paragraphs so real-world documents don't lose content.
Any other/unknown block type is skipped with a ``logger.warning`` instead of
raising, so a single malformed or unexpected block never aborts the export.

Layout modes (``export_docx(..., layout_mode=...)``):

- ``"proportional"`` (default): use each block's ``bbox`` together with the page
  size to preserve the *relative* geometry of the original scan. Images and
  tables are sized to their share of the page width (instead of being blown up
  to the full text width), blocks are horizontally aligned according to their
  position, and groups of blocks that sit side by side (overlapping vertically
  but disjoint horizontally) are laid out next to each other using invisible
  grid tables. Reconstruction is done with a recursive XY-cut over the block
  bounding boxes, so multi-column spreads no longer collapse into a single
  column. Blocks without a usable ``bbox`` fall back to full-width flow.
  Each Word sheet is sized and oriented to match its source page (aspect ratio
  preserved, clamped into the A5..A3 envelope, slim uniform margins); a new
  ``WD_SECTION.NEW_PAGE`` section is started only when the sheet size changes,
  otherwise pages are separated by plain page breaks. When ``*_middle.json``
  line data is available, per-block font sizes are estimated from the detected
  text line heights so the text scale follows the original page.
- ``"flow"``: the legacy behaviour -- every block is emitted linearly at the
  full text width, in reading order.
"""

from __future__ import annotations

import json
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt
from loguru import logger

CONTENT_LIST_SUFFIX = "_content_list.json"
MIDDLE_JSON_SUFFIX = "_middle.json"

# Block types rendered as a simple paragraph made of their "text" field.
_PLAIN_TEXT_BLOCK_TYPES = {"header", "footer", "page_number", "aside_text", "ref_text"}

# --- proportional-layout tuning constants -------------------------------------
# A vertical/horizontal gutter must be at least this fraction of the reference
# page dimension before it is treated as a real column/row separator. Kept small
# because MinerU bboxes come from layout detection and are fairly clean.
_MIN_GUTTER_FRAC = 0.006
# When deciding block alignment, a side gap smaller than this fraction of the
# region width counts as "flush" against that edge.
_ALIGN_FLUSH_FRAC = 0.06
# If a page_size is available but the bboxes overshoot it by more than this
# factor, the bboxes live in a different coordinate space than page_size, so we
# fall back to estimating the reference box from the bbox extents themselves.
_PAGE_SIZE_TOLERANCE = 1.08

# --- physical page sizing (proportional mode) ----------------------------------
# ISO A-series page dimensions in points (long side, short side). The Word sheet
# is sized to the source page, clamped into the A5..A3 envelope: anything
# smaller than A5 is stretched up to A4, anything larger than A3 is shrunk to
# A3, always preserving the source aspect ratio.
_A5_LONG_PT, _A5_SHORT_PT = 595.0, 420.0
_A4_LONG_PT, _A4_SHORT_PT = 842.0, 595.0
_A3_LONG_PT, _A3_SHORT_PT = 1191.0, 842.0
_PAGE_MARGIN_MM = 10
# Consecutive pages whose sheet dimensions differ by no more than this relative
# tolerance share one section (scanned pages jitter by a point or two).
_PAGE_DIMS_REUSE_TOL = 0.03
# Without a trustworthy page_size the bbox extents track the content, not the
# sheet; a near-square extent (long/short below this ratio) carries no reliable
# orientation signal, so such pages default to a portrait A4 sheet.
_SQUARE_ASPECT_TOLERANCE = 1.10

# --- font sizing (proportional mode) --------------------------------------------
# Word font size is estimated from middle.json text line heights: cap height is
# roughly this fraction of the detected line box height.
_FONT_LINE_FRAC = 0.7
_FONT_MIN_PT, _FONT_MAX_PT = 5.0, 14.0
# A content_list block is matched to a middle.json block when their overlap
# covers at least this fraction of the smaller box.
_FONT_MATCH_MIN_OVERLAP = 0.5
# Fallback when middle.json line data is unavailable: scale a 10pt base by the
# Word-to-source width ratio, clamped to a narrower sane range.
_FONT_FALLBACK_BASE_PT = 10.0
_FONT_FALLBACK_MIN_PT, _FONT_FALLBACK_MAX_PT = 6.0, 11.0


# --- text reflow (visual-line -> paragraph) -------------------------------------
# The hybrid backend embeds one literal "\n" per OCR-detected visual line inside
# a content_list "text" block; naive rendering turns each "\n" into a hard
# <w:br/>, so the paragraph never re-wraps in Word. reflow_text_block() splits a
# block's text back into real paragraphs by classifying every "\n" boundary as a
# BREAK (kept -- real paragraph/list-item boundary) or a JOIN (soft line-wrap,
# merged back with a space, or dehyphenated). The pipeline backend never embeds
# "\n" in "text", so this is a no-op there (single-element list, text unchanged).
#
# Design: score > 0 accumulates BREAK evidence, score < 0 accumulates JOIN
# evidence; the boundary is JOINed only when the total score is <= T_JOIN (a
# conservative default -- the OCR pipeline already placed the "\n" there, so
# ties resolve to BREAK). "fill" is a block-relative character-length proxy
# for how "full" a visual line is (hybrid has no per-line geometry):
#   fill(line) = len(line.strip()) / max(len(s.strip()) for s in segments)

_TERMINAL_CHARS = set(".!?")

# Ukrainian/Russian abbreviations whose trailing "." is not a sentence end.
_ABBR = {
    "п", "пп", "див", "табл", "рис", "ст", "стор", "буд", "вул", "м", "кв",
    "корп", "оф", "грн", "коп", "тис", "млн", "млрд", "р", "рр", "т", "ч",
    "хв", "сек", "им", "ім", "англ", "укр", "рос", "см", "ср", "напр",
    "т.д", "т.п", "т.ін",
}

# Prepositions / conjunctions that never end a wrapped line -> strong JOIN.
_PREP = {
    "в", "з", "і", "та", "й", "на", "до", "у", "від", "для", "що", "при",
    "під", "за", "из", "с", "по", "о", "об", "не", "а",
}

# Next line starts a new numbered/marked list item -> hard BREAK.
_RE_NUM = re.compile(
    r"^\s*(\d+(\.\d+)*[.)]|\(\d+\)|[а-яёіїєґ][.)]\s|[IVXLCDM]+[.)]|[-–—•●▪·*]\s)",
    re.IGNORECASE,
)
# Next line is a "Label:" / "label:" field caption.
_RE_FIELD_LABEL = re.compile(r"^\s*[\wЀ-ӿ/'\"()–— ]{1,32}:(\s|$)")
# A trailing "3.1.5" / "0,5" / "...А.3.1-5" style tail: the final "." is not
# a sentence end.
_RE_DECIMAL_TAIL = re.compile(r"\d[.,]\d*$")
_RE_HYPHEN_END = re.compile(r"[\wЀ-ӿ]-$")

# Left-hand stems of hyphenated compounds where the hyphen is a real
# orthographic hyphen (kept on JOIN), not an OCR syllable-transfer break
# (dropped on JOIN). No dictionary is available, so anything not matching one
# of these stems is treated as a soft/syllable hyphen and is removed --
# per design, an unrecognised fragment defaults to dehyphenation.
_HYPHEN_KEEP_STEMS = {
    "науково", "проектно", "монтажно", "будівельно", "ремонтно",
    "приймання", "здавання", "техніко", "інженерно", "нормативно",
    "організаційно", "санітарно", "пожежно", "вибухо", "електро",
}

_T_JOIN = -1.0
_FILL_FULL = 0.86
_FILL_MID_LO = 0.60
_SHORT_HEAD = 0.78

_W_SEMI = 3.0
_W_NUM_NEXT = 3.0
_W_NUM_HEAD_SHORT = 2.6
_W_COLON = 2.0
_W_TERM_CAP = 1.6
_W_FIELD_LABEL = 1.2
_W_TERM_LOW = 0.5
_W_HYPHEN = -3.0
_W_COMMA = -2.6
_W_PREP = -2.4
_W_WRAP_FULL = -2.4
_W_WRAP_MID = -0.8


def _last_word(s: str) -> str:
    match = re.search(r"(\S+)\s*$", s)
    return match.group(1) if match else ""


def _prev_terminal(s: str) -> str:
    """Classify the trailing token of a visual line: ``"term"`` (real sentence
    end), ``"colon"``, ``"semi"``, ``"comma"``, ``"hyphen"``, ``"prep"``, or
    ``""`` for nothing special."""
    s = s.rstrip()
    if not s:
        return ""
    if _RE_HYPHEN_END.search(s):
        return "hyphen"
    ch = s[-1]
    if ch == ";":
        return "semi"
    if ch == ":":
        return "colon"
    if ch == ",":
        return "comma"
    if ch in _TERMINAL_CHARS:
        if ch == "." and _RE_DECIMAL_TAIL.search(s):
            return ""  # decimal/version/section-number tail, not a sentence end
        word = re.sub(r"^[^\wЀ-ӿ]+", "", _last_word(s).rstrip(".").lower())
        if ch == "." and (word in _ABBR or (len(word) == 1 and word.isalpha())):
            return ""
        return "term"
    last = re.sub(r"[^\wЀ-ӿ]", "", _last_word(s).lower())
    if last in _PREP:
        return "prep"
    return ""


def _starts_lower(s: str) -> bool:
    s = s.lstrip()
    return bool(s) and s[0].islower()


def _keep_hyphen(prev: str) -> bool:
    """Whether the trailing ``-`` in ``prev`` is a real compound-word hyphen
    that must survive a JOIN (e.g. ``"науково-технічний"``), as opposed to an
    OCR line-wrap/syllable-transfer hyphen that gets dropped on JOIN (e.g.
    ``"відси-" + "паного"`` -> ``"відсипаного"``).

    Rule actually applied (see module docstring above for rationale): kept
    only when the word fragment immediately preceding the hyphen equals, or
    ends with, one of ``_HYPHEN_KEEP_STEMS``. There is no dictionary to
    validate arbitrary fragments as "real words", so anything outside the
    whitelist is dehyphenated -- this matches the documented fallback ("if in
    doubt, rely on the whitelist") and is what keeps a plain syllable break
    like "відси-/паного" from being misdetected as a compound.
    """
    stripped = prev.rstrip()
    match = re.search(r"([\wЀ-ӿ]+)-$", stripped)
    if not match:
        return False
    left = match.group(1).lower()
    return any(left == stem or left.endswith(stem) for stem in _HYPHEN_KEEP_STEMS)


def _classify_boundary(prev: str, nxt: str, fill_prev: float, prev_is_num: bool) -> str:
    """Score one "\\n" boundary between two visual lines and return ``"JOIN"``
    or ``"BREAK"``."""
    score = 0.0
    prev_end = _prev_terminal(prev)
    low = _starts_lower(nxt)
    num_next = bool(_RE_NUM.match(nxt))
    field = bool(_RE_FIELD_LABEL.match(nxt))

    # --- BREAK evidence ---
    if prev_end == "semi":
        score += _W_SEMI
    if num_next:
        score += _W_NUM_NEXT
    if prev_is_num and fill_prev < _SHORT_HEAD:
        score += _W_NUM_HEAD_SHORT
    if prev_end == "colon":
        score += _W_COLON
    if prev_end == "term":
        score += _W_TERM_LOW if low else _W_TERM_CAP
    if field and not num_next:
        score += _W_FIELD_LABEL

    # --- JOIN evidence ---
    if prev_end == "hyphen":
        score += _W_HYPHEN
    if prev_end == "comma":
        score += _W_COMMA
    if prev_end == "prep":
        score += _W_PREP
    if prev_end in ("", "prep", "comma") and low:
        if fill_prev >= _FILL_FULL:
            score += _W_WRAP_FULL
        elif fill_prev >= _FILL_MID_LO:
            score += _W_WRAP_MID

    return "JOIN" if score <= _T_JOIN else "BREAK"


def reflow_text_block(text: str) -> list[str]:
    """Split a content_list block's ``text`` field into a list of paragraphs.

    Every embedded ``"\\n"`` boundary is classified as BREAK (a hard boundary
    -- kept as a separate paragraph, e.g. list items, sentence ends) or JOIN
    (a soft line-wrap -- merged back into the previous paragraph with a space,
    or dehyphenated, see :func:`_keep_hyphen`). Text without any ``"\\n"``
    (the common pipeline-backend case) is returned unchanged as a
    single-element list.
    """
    if not text:
        return [text]
    segments = text.split("\n")
    if len(segments) == 1:
        return [text]

    lengths = [len(seg.strip()) for seg in segments]
    max_len = max(lengths) if lengths else 1

    paragraphs: list[str] = []
    current = segments[0]
    for i in range(len(segments) - 1):
        prev, nxt = segments[i], segments[i + 1]
        fill = (lengths[i] / max_len) if max_len else 0.0
        prev_is_num = bool(_RE_NUM.match(prev))
        decision = _classify_boundary(prev, nxt, fill, prev_is_num)
        if decision == "JOIN":
            if _RE_HYPHEN_END.search(prev.rstrip()):
                if _keep_hyphen(prev):
                    current = current.rstrip() + nxt.lstrip()
                else:
                    current = current.rstrip()[:-1] + nxt.lstrip()
            else:
                current = current.rstrip() + " " + nxt.lstrip()
        else:
            paragraphs.append(current)
            current = nxt
    paragraphs.append(current)
    return paragraphs


class _HtmlTableParser(HTMLParser):
    """Minimal HTML table parser: extracts rows of cells with text/rowspan/colspan.

    Only understands the subset of HTML MinerU emits for ``table_body``
    (``<table>``, ``<tr>``, ``<td>``/``<th>``, ``<br>``); anything else is
    treated as plain inline text inside the current cell.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.rows: list[list[dict[str, Any]]] = []
        self._current_row: list[dict[str, Any]] | None = None
        self._in_cell = False
        self._cell_text_parts: list[str] = []
        self._cell_attrs: dict[str, str] = {}

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = {key: value for key, value in attrs if value is not None}
        if tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._in_cell = True
            self._cell_text_parts = []
            self._cell_attrs = attrs_dict
        elif tag == "br" and self._in_cell:
            self._cell_text_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in ("td", "th"):
            if self._current_row is None:
                self._current_row = []
            text = "".join(self._cell_text_parts).strip()
            rowspan = _safe_int(self._cell_attrs.get("rowspan"), default=1)
            colspan = _safe_int(self._cell_attrs.get("colspan"), default=1)
            self._current_row.append(
                {"text": text, "rowspan": max(1, rowspan), "colspan": max(1, colspan)}
            )
            self._in_cell = False
        elif tag == "tr":
            if self._current_row is not None:
                self.rows.append(self._current_row)
            self._current_row = None

    def handle_data(self, data: str) -> None:
        if self._in_cell:
            self._cell_text_parts.append(data)


def _safe_int(value: str | None, default: int) -> int:
    if value is None:
        return default
    try:
        return int(str(value).strip())
    except (TypeError, ValueError):
        return default


def _parse_html_table(table_html: str) -> list[list[dict[str, Any]]]:
    parser = _HtmlTableParser()
    parser.feed(table_html or "")
    return parser.rows


def _build_table_grid(
    rows: list[list[dict[str, Any]]],
) -> tuple[dict[tuple[int, int], dict[str, Any]], int, int]:
    """Place each cell at its top-left anchor position in a rectangular grid.

    Returns a mapping of (row, col) -> cell spec (only anchor positions are
    present) plus the total number of rows/cols in the resulting grid.
    """
    grid: dict[tuple[int, int], dict[str, Any]] = {}
    occupied: set[tuple[int, int]] = set()
    max_cols = 0

    for row_idx, row in enumerate(rows):
        col_idx = 0
        for cell in row:
            while (row_idx, col_idx) in occupied:
                col_idx += 1
            grid[(row_idx, col_idx)] = cell
            for delta_row in range(cell["rowspan"]):
                for delta_col in range(cell["colspan"]):
                    occupied.add((row_idx + delta_row, col_idx + delta_col))
            col_idx += cell["colspan"]
        max_cols = max(max_cols, col_idx)

    return grid, len(rows), max_cols


def _set_cell_text(cell, text: str, font_pt: float | None = None) -> None:
    lines = text.split("\n") if text else [""]
    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)
    if font_pt:
        for paragraph in cell.paragraphs:
            _apply_font_size(paragraph, font_pt)


def _apply_font_size(paragraph, font_pt: float | None) -> None:
    """Apply the estimated font size plus a dense vertical rhythm (single line
    spacing, minimal space after) so proportional pages pack like the source
    scan. No-op when ``font_pt`` is None (e.g. flow mode)."""
    if not font_pt:
        return
    for run in paragraph.runs:
        run.font.size = Pt(font_pt)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0


# --- container abstraction ----------------------------------------------------
# Both ``Document`` and table ``_Cell`` expose ``add_paragraph`` / ``add_table``
# with the same signature, so most helpers accept either as ``parent``. Pictures
# need to go through a run, and headings are added as styled paragraphs so they
# work inside cells too.


def _add_paragraph(parent, text: str | None = None, style: str | None = None):
    paragraph = parent.add_paragraph(style=style) if style else parent.add_paragraph()
    if text:
        paragraph.add_run(text)
    return paragraph


# --- heading sizing/spacing (vertical-compression tuning) ----------------------
# Heading run size, when a page font estimate is available, is
# font_pt * level coefficient (never below font_pt -- headings must never
# read smaller than body text). Coefficients decrease with depth so H1 stays
# visually dominant over H3 while both still track the page's overall scale
# instead of a fixed style size that can dwarf a small proportional page.
_HEADING_LEVEL_COEF = {1: 1.4, 2: 1.3, 3: 1.2}
_HEADING_LEVEL_COEF_DEFAULT = 1.15
# Spacing around headings, tuned empirically against a real 49-page hybrid
# document (see export verification report): the ~10-24pt style defaults for
# Heading 1-3 turned out to be the single largest contributor to the extra
# page count (measured -2 rendered pages from this change alone, vs a
# negligible effect from table cell margins on that document). Readability is
# preserved through the heading's larger/styled run (see
# ``_HEADING_LEVEL_COEF``, bold/color from the "Heading N" style) rather than
# surrounding whitespace, so zero space_before/space_after is acceptable here.
_HEADING_SPACE_BEFORE_PT = 0.0
_HEADING_SPACE_AFTER_PT = 0.0


def _add_heading_paragraph(parent, text: str, level: int, font_pt: float | None = None):
    """Add a heading paragraph using the built-in "Heading N" style, then
    tighten its vertical footprint for page-count parity with the source
    scan: minimal space_before/space_after (see ``_HEADING_SPACE_*``), and --
    when a page font-size estimate is available (proportional mode) -- the
    run size is rescaled from the style's fixed default to ``font_pt`` scaled
    up by a per-level coefficient (see ``_HEADING_LEVEL_COEF``), so headings
    stay larger than body text but track the page's overall font scale
    instead of dominating a small proportional page. ``font_pt`` is None in
    flow mode / when no estimate is available, in which case the style's
    default run size is left untouched (only spacing is tightened)."""
    level = max(1, min(int(level), 9))
    try:
        paragraph = parent.add_paragraph(text, style=f"Heading {level}")
    except KeyError:
        # Style missing from the template: degrade gracefully to bold text.
        paragraph = parent.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True

    if font_pt:
        coef = _HEADING_LEVEL_COEF.get(level, _HEADING_LEVEL_COEF_DEFAULT)
        heading_pt = max(font_pt * coef, font_pt)  # never below body text size
        for run in paragraph.runs:
            run.font.size = Pt(heading_pt)

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(_HEADING_SPACE_BEFORE_PT)
    fmt.space_after = Pt(_HEADING_SPACE_AFTER_PT)
    return paragraph


_LABEL_VALUE_MAX_COLS = 2
_LABEL_VALUE_MAX_COLSPAN = 2
_LABEL_VALUE_MEDIAN_LEN = 40
_LABEL_VALUE_MAX_NUMERIC_FRAC = 0.5
_LABEL_VALUE_MIN_COLON_FRAC = 0.6
_RE_NUMERIC_CELL = re.compile(r"^[\d.,\-–—\s%]+$")
_RE_ENDS_COLON = re.compile(r":\s*$")


def _is_label_value_table(rows: list[list[dict[str, Any]]]) -> bool:
    """Heuristic: does this parsed HTML table look like a "label: value"
    pseudo-table (e.g. a document front-matter block such as "РОЗРОБЛЕНО:" /
    "ЗАТВЕРДЖЕНО:" rows) rather than a real data grid?

    A table qualifies -- and should therefore render *without* visible
    borders (like the invisible layout tables used for side-by-side blocks,
    see :func:`_remove_table_borders`) -- only when ALL of the following
    hold:

    - every row has at most ``_LABEL_VALUE_MAX_COLS`` (2) cells: a label
      column plus an optional value column;
    - no cell spans more than ``_LABEL_VALUE_MAX_COLSPAN`` (2) grid columns,
      which would indicate a real multi-column header/merge;
    - the left-hand column reads as short text labels: the median cell text
      length is below ``_LABEL_VALUE_MEDIAN_LEN`` (~40) characters;
    - the left-hand column is not a numeric grid (e.g. a row-number column
      in a real data table): at most half its cells are purely
      numeric/decimal/percentage tokens;
    - the left-hand column actually reads as field captions, not merely
      short/sparse cells: at least ``_LABEL_VALUE_MIN_COLON_FRAC`` (60%) of
      its non-empty cells end with a colon (e.g. "РОЗРОБЛЕНО:",
      "РОЗРОБНИКИ:"). This extra check is what tells a real front-matter
      block apart from a genuinely empty/sparse data table (e.g. a blank
      journal/log grid) that happens to be narrow and short-celled but whose
      cells are not colon-terminated captions.

    Real data grids -- multi-column tables, tables with wide merged headers,
    tables whose first column is a long caption or a numbering column, or
    tables whose left column is short/sparse but not colon-terminated
    captions (e.g. an empty journal table) -- return ``False`` and keep the
    regular "Table Grid" borders. Empty input also returns ``False``.
    """
    if not rows:
        return False
    if any(len(row) > _LABEL_VALUE_MAX_COLS for row in rows):
        return False
    if any(cell.get("colspan", 1) > _LABEL_VALUE_MAX_COLSPAN for row in rows for cell in row):
        return False

    left_texts = [row[0].get("text", "") or "" for row in rows if row]
    if not left_texts:
        return False

    lengths = sorted(len(t) for t in left_texts)
    median_len = lengths[len(lengths) // 2]
    if median_len >= _LABEL_VALUE_MEDIAN_LEN:
        return False

    numeric = sum(1 for t in left_texts if t.strip() and _RE_NUMERIC_CELL.match(t.strip()))
    if (numeric / len(left_texts)) > _LABEL_VALUE_MAX_NUMERIC_FRAC:
        return False

    non_empty = [t.strip() for t in left_texts if t.strip()]
    if not non_empty:
        return False
    colon_terminated = sum(1 for t in non_empty if _RE_ENDS_COLON.search(t))
    if (colon_terminated / len(non_empty)) < _LABEL_VALUE_MIN_COLON_FRAC:
        return False

    return True


def _tighten_table_cell_margins(table, margin_pt: float = 0.75) -> None:
    """Shrink Word's default Table Grid cell padding (~0.08in per side) to
    ``margin_pt``. Default cell margins add several points of dead space per
    cell, compounding across every row of a long real data grid -- a
    meaningful share of the extra page count vs. the source scan."""
    dxa = str(max(0, int(margin_pt * 20)))  # dxa = twentieths of a point
    tbl_pr = table._element.tblPr
    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = tbl_pr.makeelement(qn("w:tblCellMar"), {})
        tbl_pr.append(cell_mar)
    for edge in ("top", "left", "bottom", "right"):
        element = cell_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = cell_mar.makeelement(qn(f"w:{edge}"), {})
            cell_mar.append(element)
        element.set(qn("w:type"), "dxa")
        element.set(qn("w:w"), dxa)


def _remove_table_borders(table) -> None:
    """Make a layout table invisible (no borders) so side-by-side blocks read
    as free-floating regions rather than a visible grid."""
    tbl_pr = table._element.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_fixed_table_layout(table) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._element.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = tbl_pr.makeelement(qn("w:tblLayout"), {})
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_column_widths(table, col_widths_emu: list[int]) -> None:
    """Force fixed column widths (in EMU). ``col_widths_emu`` length must match
    the number of grid columns."""
    _set_fixed_table_layout(table)
    total = sum(col_widths_emu)
    if total > 0:
        table.width = Emu(total)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(col_widths_emu):
                try:
                    cell.width = Emu(col_widths_emu[idx])
                except Exception:  # pragma: no cover - defensive
                    pass


# --- geometry helpers ---------------------------------------------------------


def _block_bbox(block: dict[str, Any]) -> tuple[float, float, float, float] | None:
    bbox = block.get("bbox")
    if not isinstance(bbox, (list, tuple)) or len(bbox) < 4:
        return None
    try:
        x0, y0, x1, y1 = (float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
    except (TypeError, ValueError):
        return None
    if x1 < x0:
        x0, x1 = x1, x0
    if y1 < y0:
        y0, y1 = y1, y0
    return x0, y0, x1, y1


def _page_reference_box(
    blocks: list[dict[str, Any]], page_size: tuple[float, float] | None
) -> tuple[float, float, float, float]:
    """Return the (x0, y0, x1, y1) reference box the proportions are computed
    against.

    Prefers ``page_size`` when the block bboxes actually fit inside it (pipeline
    backend), otherwise estimates the box from the bbox extents (e.g. the hybrid
    backend emits content_list bboxes in a different, larger coordinate space
    than middle.json's page_size).
    """
    boxes = [b for b in (_block_bbox(bl) for bl in blocks) if b is not None]
    if not boxes:
        if page_size:
            return 0.0, 0.0, float(page_size[0]), float(page_size[1])
        return 0.0, 0.0, 1.0, 1.0
    ext_x1 = max(b[2] for b in boxes)
    ext_y1 = max(b[3] for b in boxes)
    if page_size:
        pw, ph = float(page_size[0]), float(page_size[1])
        if pw > 0 and ph > 0 and ext_x1 <= pw * _PAGE_SIZE_TOLERANCE and ext_y1 <= ph * _PAGE_SIZE_TOLERANCE:
            return 0.0, 0.0, pw, ph
    return 0.0, 0.0, ext_x1, ext_y1


def _physical_page_pts(src_width: float, src_height: float) -> tuple[float, float]:
    """Map a source page size (points) to the Word sheet size (points).

    The source aspect ratio is always preserved; the size is clamped into the
    A5..A3 envelope (smaller than A5 -> stretched to A4, larger than A3 ->
    shrunk to A3). Degenerate/unknown sizes fall back to portrait A4.
    """
    if src_width <= 2 or src_height <= 2:  # unit/degenerate reference box
        return _A4_SHORT_PT, _A4_LONG_PT
    long_side = max(src_width, src_height)
    short_side = min(src_width, src_height)
    scale = 1.0
    if long_side < _A5_LONG_PT and short_side < _A5_SHORT_PT:
        scale = min(_A4_LONG_PT / long_side, _A4_SHORT_PT / short_side)
    elif long_side > _A3_LONG_PT or short_side > _A3_SHORT_PT:
        scale = min(_A3_LONG_PT / long_side, _A3_SHORT_PT / short_side)
    return src_width * scale, src_height * scale


def _configure_page_section(section, page_w_pt: float, page_h_pt: float) -> None:
    """Size a section to the given sheet dimensions (points), set the matching
    orientation flag and apply the slim uniform page margins."""
    section.page_width = Pt(page_w_pt)
    section.page_height = Pt(page_h_pt)
    section.orientation = WD_ORIENT.LANDSCAPE if page_w_pt > page_h_pt else WD_ORIENT.PORTRAIT
    margin = Mm(_PAGE_MARGIN_MM)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin


def _section_usable_width(section) -> int:
    """Usable content width (EMU) of a section: page width minus side margins."""
    return int(section.page_width - section.left_margin - section.right_margin)


# --- font size estimation -------------------------------------------------------


def _median(values: list[float]) -> float:
    ordered = sorted(values)
    return ordered[len(ordered) // 2]


def _clamp(value: float, lo: float, hi: float) -> float:
    return max(lo, min(hi, value))


def _collect_text_line_heights(block: dict[str, Any]) -> list[float]:
    """Collect text line heights (middle.json units) from a middle.json block
    subtree. ``*_body`` leaves (image/table bodies) are skipped: their "lines"
    span the whole body and say nothing about the text size."""
    heights: list[float] = []

    def walk(node: dict[str, Any]) -> None:
        node_type = str(node.get("type") or "")
        if not node_type.endswith("_body"):
            for line in node.get("lines") or []:
                if not isinstance(line, dict):
                    continue
                lb = line.get("bbox")
                if isinstance(lb, (list, tuple)) and len(lb) >= 4:
                    try:
                        height = float(lb[3]) - float(lb[1])
                    except (TypeError, ValueError):
                        continue
                    if height > 0:
                        heights.append(height)
        for sub in node.get("blocks") or []:
            if isinstance(sub, dict):
                walk(sub)

    walk(block)
    return heights


def _overlap_frac(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    """Intersection area over the smaller box's area (0 when disjoint)."""
    ix = max(0.0, min(a[2], b[2]) - max(a[0], b[0]))
    iy = max(0.0, min(a[3], b[3]) - max(a[1], b[1]))
    inter = ix * iy
    if inter <= 0:
        return 0.0
    area_a = max(0.0, (a[2] - a[0]) * (a[3] - a[1]))
    area_b = max(0.0, (b[2] - b[0]) * (b[3] - b[1]))
    smaller = min(area_a, area_b)
    return inter / smaller if smaller > 0 else 0.0


def _build_page_font_map(
    page_blocks: list[dict[str, Any]],
    region: tuple[float, float, float, float],
    middle_page: dict[str, Any] | None,
    usable_width_emu: int,
) -> tuple[dict[int, float], float | None]:
    """Estimate per-block and page-default font sizes (pt) for one page.

    Line heights come from middle.json; they are converted to Word points via
    the width ratio between the Word text column and the middle.json page (the
    sheet keeps the source aspect ratio, so the width ratio is the uniform
    scale factor). Returns ``(font_by_block_id, page_default_pt)``.
    """
    usable_width_pt = usable_width_emu / 12700.0
    ref_w = region[2] - region[0]
    ref_h = region[3] - region[1]

    font_map: dict[int, float] = {}
    page_default: float | None = None

    mid_size = (middle_page or {}).get("page_size")
    mid_blocks = (middle_page or {}).get("blocks") or []
    if (
        isinstance(mid_size, (list, tuple))
        and len(mid_size) >= 2
        and float(mid_size[0]) > 0
        and float(mid_size[1]) > 0
        and mid_blocks
    ):
        mid_w, mid_h = float(mid_size[0]), float(mid_size[1])
        scale_pt = usable_width_pt / mid_w

        page_heights = [h for mb in mid_blocks for h in mb.get("heights") or []]
        if page_heights:
            page_default = _clamp(
                _FONT_LINE_FRAC * _median(page_heights) * scale_pt, _FONT_MIN_PT, _FONT_MAX_PT
            )

        if ref_w > 0 and ref_h > 0:
            sx, sy = mid_w / ref_w, mid_h / ref_h
            for block in page_blocks:
                bbox = _block_bbox(block)
                if bbox is None:
                    continue
                scaled = (bbox[0] * sx, bbox[1] * sy, bbox[2] * sx, bbox[3] * sy)
                best_heights: list[float] | None = None
                best_overlap = _FONT_MATCH_MIN_OVERLAP
                for mb in mid_blocks:
                    mb_bbox = mb.get("bbox")
                    if mb_bbox is None:
                        continue
                    overlap = _overlap_frac(scaled, mb_bbox)
                    if overlap >= best_overlap and mb.get("heights"):
                        best_overlap = overlap
                        best_heights = mb["heights"]
                if best_heights:
                    font_map[id(block)] = _clamp(
                        _FONT_LINE_FRAC * _median(best_heights) * scale_pt,
                        _FONT_MIN_PT,
                        _FONT_MAX_PT,
                    )

    if page_default is None and ref_w > 0:
        # No usable middle.json data: scale a 10pt base font by how much the
        # source page (treated as points) is squeezed into the Word column.
        page_default = _clamp(
            _FONT_FALLBACK_BASE_PT * (usable_width_pt / ref_w),
            _FONT_FALLBACK_MIN_PT,
            _FONT_FALLBACK_MAX_PT,
        )

    return font_map, page_default


def _cut_1d(intervals: list[tuple[float, float]], lo: float, hi: float, min_gap: float) -> list[tuple[float, float]]:
    """Given 1-D intervals within [lo, hi], return the maximal covered clusters
    (merging intervals separated by a gap smaller than ``min_gap``). Returns a
    list of (start, end) cluster spans; length > 1 means a real cut exists."""
    if not intervals:
        return []
    ordered = sorted(intervals)
    clusters: list[list[float]] = [[ordered[0][0], ordered[0][1]]]
    for start, end in ordered[1:]:
        if start - clusters[-1][1] >= min_gap:
            clusters.append([start, end])
        else:
            clusters[-1][1] = max(clusters[-1][1], end)
    return [(c[0], c[1]) for c in clusters]


def _horizontal_align(x0: float, x1: float, rx0: float, rx1: float):
    region_w = rx1 - rx0
    if region_w <= 0:
        return WD_ALIGN_PARAGRAPH.LEFT
    left_gap = (x0 - rx0) / region_w
    right_gap = (rx1 - x1) / region_w
    if left_gap <= _ALIGN_FLUSH_FRAC and right_gap <= _ALIGN_FLUSH_FRAC:
        return WD_ALIGN_PARAGRAPH.LEFT  # fills the region
    if left_gap <= _ALIGN_FLUSH_FRAC:
        return WD_ALIGN_PARAGRAPH.LEFT
    if right_gap <= _ALIGN_FLUSH_FRAC:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if abs(left_gap - right_gap) <= _ALIGN_FLUSH_FRAC:
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT if left_gap < right_gap else WD_ALIGN_PARAGRAPH.RIGHT


_TABLE_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: WD_TABLE_ALIGNMENT.LEFT,
    WD_ALIGN_PARAGRAPH.CENTER: WD_TABLE_ALIGNMENT.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT: WD_TABLE_ALIGNMENT.RIGHT,
}


# --- block renderers ----------------------------------------------------------


def _resolve_image_path(images_dir: Path, img_path: str) -> Path | None:
    img_path = str(img_path or "").strip()
    if not img_path:
        return None

    candidates = [
        images_dir.parent / img_path,
        images_dir / Path(img_path).name,
        Path(img_path),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return None


def _render_table_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    font_pt = ctx.font_for(block)
    for caption in block.get("table_caption") or []:
        if caption:
            paragraph = _add_paragraph(parent, caption)
            paragraph.runs[0].italic = True
            _apply_font_size(paragraph, font_pt)

    rows = _parse_html_table(block.get("table_body") or "")
    if not rows:
        logger.warning("Table block has no parseable rows, skipping table body.")
    else:
        grid, n_rows, n_cols = _build_table_grid(rows)
        if n_rows > 0 and n_cols > 0:
            table = parent.add_table(rows=n_rows, cols=n_cols)
            if _is_label_value_table(rows):
                # "Label: value" pseudo-table (e.g. front-matter blocks like
                # "РОЗРОБЛЕНО:" / "ЗАТВЕРДЖЕНО:") -- not a real data grid, so
                # render it borderless like the invisible layout tables.
                _remove_table_borders(table)
            else:
                table.style = "Table Grid"
                _tighten_table_cell_margins(table)

            # Merge spanning cells first, then fill in text, so text ends up
            # on the already-merged (top-left) cell.
            for (row_idx, col_idx), cell_spec in grid.items():
                row_span = cell_spec["rowspan"]
                col_span = cell_spec["colspan"]
                if row_span == 1 and col_span == 1:
                    continue
                end_row = min(row_idx + row_span - 1, n_rows - 1)
                end_col = min(col_idx + col_span - 1, n_cols - 1)
                if end_row == row_idx and end_col == col_idx:
                    continue
                try:
                    table.cell(row_idx, col_idx).merge(table.cell(end_row, end_col))
                except Exception as exc:  # pragma: no cover - defensive
                    logger.warning(f"Failed to merge table cell ({row_idx},{col_idx}): {exc}")

            for (row_idx, col_idx), cell_spec in grid.items():
                try:
                    _set_cell_text(table.cell(row_idx, col_idx), cell_spec["text"], font_pt)
                except Exception as exc:  # pragma: no cover - defensive
                    logger.warning(f"Failed to set table cell ({row_idx},{col_idx}) text: {exc}")

            # Proportional width: scale the table to its share of the region and
            # distribute equally across the grid columns.
            bbox = _block_bbox(block)
            if ctx.proportional and bbox is not None and ctx.region_width > 0:
                frac = min(1.0, (bbox[2] - bbox[0]) / ctx.region_width)
                target = max(1, int(ctx.available_width * frac))
                per_col = max(1, target // n_cols)
                _set_table_column_widths(table, [per_col] * n_cols)
                align = _horizontal_align(bbox[0], bbox[2], ctx.region_x0, ctx.region_x1)
                table.alignment = _TABLE_ALIGN_MAP.get(align, WD_TABLE_ALIGNMENT.LEFT)

    for footnote in block.get("table_footnote") or []:
        if footnote:
            paragraph = _add_paragraph(parent, footnote)
            paragraph.runs[0].italic = True
            _apply_font_size(paragraph, font_pt)


def _render_image_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    font_pt = ctx.font_for(block)
    for caption in block.get("image_caption") or []:
        if caption:
            paragraph = _add_paragraph(parent, caption)
            paragraph.runs[0].italic = True
            _apply_font_size(paragraph, font_pt)

    img_path = block.get("img_path")
    resolved_path = _resolve_image_path(ctx.images_dir, img_path) if img_path else None
    if resolved_path is None:
        logger.warning(f"Image not found for content_list block, skipping: {img_path!r}")
    else:
        bbox = _block_bbox(block)
        if ctx.proportional and bbox is not None and ctx.region_width > 0:
            frac = min(1.0, (bbox[2] - bbox[0]) / ctx.region_width)
            width = max(1, int(ctx.available_width * frac))
            align = _horizontal_align(bbox[0], bbox[2], ctx.region_x0, ctx.region_x1)
        else:
            width = ctx.available_width
            align = WD_ALIGN_PARAGRAPH.LEFT
        paragraph = parent.add_paragraph()
        paragraph.alignment = align
        run = paragraph.add_run()
        try:
            run.add_picture(str(resolved_path), width=Emu(width))
        except Exception as exc:
            logger.warning(f"Failed to embed image {resolved_path}: {exc}")

    for footnote in block.get("image_footnote") or []:
        if footnote:
            paragraph = _add_paragraph(parent, footnote)
            paragraph.runs[0].italic = True
            _apply_font_size(paragraph, font_pt)


def _render_equation_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    text = block.get("text") or ""
    if not text.strip():
        return
    paragraph = parent.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    _apply_font_size(paragraph, ctx.font_for(block))


def _render_text_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    text = block.get("text") or ""
    if not text.strip():
        return
    text_level = block.get("text_level")
    if isinstance(text_level, int) and text_level >= 1:
        # Headings keep the style-driven structure (Heading N), but the run
        # size/spacing are tightened -- see _add_heading_paragraph.
        _add_heading_paragraph(parent, text, text_level, ctx.font_for(block))
    else:
        font_pt = ctx.font_for(block)
        paragraphs = reflow_text_block(text)
        for para_text in paragraphs:
            paragraph = _add_paragraph(parent, para_text)
            _apply_font_size(paragraph, font_pt)


def _render_list_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    font_pt = ctx.font_for(block)
    for item in block.get("list_items") or []:
        if item:
            paragraph = _add_paragraph(parent, str(item), style="List Bullet")
            _apply_font_size(paragraph, font_pt)


def _render_plain_text_block(parent, block: dict[str, Any], ctx: "_LayoutCtx") -> None:
    text = block.get("text") or ""
    if text.strip():
        paragraph = _add_paragraph(parent, text)
        _apply_font_size(paragraph, ctx.font_for(block))


def _render_block(parent, block: dict[str, Any], ctx: "_LayoutCtx", index: int) -> None:
    if not isinstance(block, dict):
        logger.warning(f"Skipping non-dict content_list entry at index {index}: {block!r}")
        return
    block_type = block.get("type")
    try:
        if block_type == "text":
            _render_text_block(parent, block, ctx)
        elif block_type == "table":
            _render_table_block(parent, block, ctx)
        elif block_type == "image":
            _render_image_block(parent, block, ctx)
        elif block_type == "equation":
            _render_equation_block(parent, block, ctx)
        elif block_type == "list":
            _render_list_block(parent, block, ctx)
        elif block_type in _PLAIN_TEXT_BLOCK_TYPES:
            _render_plain_text_block(parent, block, ctx)
        else:
            logger.warning(
                f"Unsupported content_list block type at index {index}: {block_type!r}, skipping."
            )
    except Exception as exc:
        logger.warning(f"Failed to render content_list block at index {index} ({block_type!r}): {exc}")


# --- proportional layout (recursive XY-cut) -----------------------------------


class _LayoutCtx:
    """Per-render context passed down the recursion.

    ``region_x0``/``region_x1`` describe the horizontal span (in bbox units) of
    the region currently being rendered, ``region_width`` its width, and
    ``available_width`` the matching width in EMU inside the current container.
    """

    __slots__ = (
        "images_dir",
        "proportional",
        "region_x0",
        "region_x1",
        "available_width",
        "font_map",
        "default_font_pt",
    )

    def __init__(self, images_dir: Path, proportional: bool) -> None:
        self.images_dir = images_dir
        self.proportional = proportional
        self.region_x0 = 0.0
        self.region_x1 = 1.0
        self.available_width = 0
        # Per-page font estimates: id(block) -> pt, plus the page-wide default.
        self.font_map: dict[int, float] = {}
        self.default_font_pt: float | None = None

    @property
    def region_width(self) -> float:
        return self.region_x1 - self.region_x0

    def font_for(self, block: dict[str, Any]) -> float | None:
        """Estimated font size (pt) for a block, or None to keep style default."""
        return self.font_map.get(id(block), self.default_font_pt)

    def clone_for(self, x0: float, x1: float, available_width: int) -> "_LayoutCtx":
        ctx = _LayoutCtx(self.images_dir, self.proportional)
        ctx.region_x0 = x0
        ctx.region_x1 = x1
        ctx.available_width = available_width
        ctx.font_map = self.font_map
        ctx.default_font_pt = self.default_font_pt
        return ctx


def _render_region(
    parent,
    blocks: list[dict[str, Any]],
    region: tuple[float, float, float, float],
    available_width: int,
    ctx: _LayoutCtx,
    prefer_vertical: bool,
    depth: int,
) -> None:
    """Recursively lay out ``blocks`` inside ``region`` using an XY-cut.

    Vertical cuts (columns) become invisible grid tables; horizontal cuts (rows)
    are rendered sequentially into the same container. ``prefer_vertical`` picks
    which axis to try first and is flipped on each level so the cut direction
    alternates like a classic XY-cut.
    """
    rx0, ry0, rx1, ry1 = region
    region_w = rx1 - rx0
    region_h = ry1 - ry0

    boxed = [(b, _block_bbox(b)) for b in blocks]
    boxed = [(b, bb) for b, bb in boxed if bb is not None]
    if not boxed:
        return

    if len(boxed) == 1 or depth > 12:
        local = ctx.clone_for(rx0, rx1, available_width)
        for b, _bb in sorted(boxed, key=lambda pair: (pair[1][1], pair[1][0])):
            _render_block(parent, b, local, 0)
        return

    min_gap_x = max(region_w * _MIN_GUTTER_FRAC, 1.0)
    min_gap_y = max(region_h * _MIN_GUTTER_FRAC, 1.0)

    def try_vertical() -> list[tuple[float, float]]:
        spans = _cut_1d([(bb[0], bb[2]) for _b, bb in boxed], rx0, rx1, min_gap_x)
        return spans if len(spans) > 1 else []

    def try_horizontal() -> list[tuple[float, float]]:
        spans = _cut_1d([(bb[1], bb[3]) for _b, bb in boxed], ry0, ry1, min_gap_y)
        return spans if len(spans) > 1 else []

    order = (try_vertical, try_horizontal) if prefer_vertical else (try_horizontal, try_vertical)
    for cut in order:
        spans = cut()
        if not spans:
            continue
        if cut is try_vertical:
            _render_columns(parent, boxed, spans, region, available_width, ctx, depth)
        else:
            _render_rows(parent, boxed, spans, region, available_width, ctx, depth)
        return

    # No further cut possible: render remaining blocks as flow, top-to-bottom.
    local = ctx.clone_for(rx0, rx1, available_width)
    for b, _bb in sorted(boxed, key=lambda pair: (pair[1][1], pair[1][0])):
        _render_block(parent, b, local, 0)


def _render_columns(parent, boxed, spans, region, available_width, ctx, depth) -> None:
    rx0, ry0, rx1, ry1 = region
    region_w = rx1 - rx0 or 1.0
    # Assign each block to the column whose span contains its horizontal centre.
    groups: list[list[tuple[dict[str, Any], tuple[float, float, float, float]]]] = [[] for _ in spans]
    for b, bb in boxed:
        centre = (bb[0] + bb[2]) / 2.0
        idx = 0
        for i, (s0, s1) in enumerate(spans):
            if s0 <= centre <= s1:
                idx = i
                break
        else:
            # centre outside every span: snap to nearest
            idx = min(range(len(spans)), key=lambda i: abs(centre - (spans[i][0] + spans[i][1]) / 2.0))
        groups[idx].append((b, bb))

    col_widths = [max(1, int(available_width * ((s1 - s0) / region_w))) for s0, s1 in spans]
    table = parent.add_table(rows=1, cols=len(spans))
    _remove_table_borders(table)
    _set_table_column_widths(table, col_widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (s0, s1) in enumerate(spans):
        cell = table.cell(0, i)
        # A fresh cell already has one empty paragraph; drop it to avoid a blank
        # leading line inside the column.
        _clear_cell(cell)
        col_region = (s0, ry0, s1, ry1)
        _render_region(cell, [b for b, _bb in groups[i]], col_region, col_widths[i], ctx, prefer_vertical=False, depth=depth + 1)


def _render_rows(parent, boxed, spans, region, available_width, ctx, depth) -> None:
    rx0, ry0, rx1, ry1 = region
    groups: list[list[tuple[dict[str, Any], tuple[float, float, float, float]]]] = [[] for _ in spans]
    for b, bb in boxed:
        centre = (bb[1] + bb[3]) / 2.0
        idx = 0
        for i, (s0, s1) in enumerate(spans):
            if s0 <= centre <= s1:
                idx = i
                break
        else:
            idx = min(range(len(spans)), key=lambda i: abs(centre - (spans[i][0] + spans[i][1]) / 2.0))
        groups[idx].append((b, bb))

    for i, (s0, s1) in enumerate(spans):
        row_region = (rx0, s0, rx1, s1)
        _render_region(parent, [b for b, _bb in groups[i]], row_region, available_width, ctx, prefer_vertical=True, depth=depth + 1)


def _clear_cell(cell) -> None:
    """Remove the default empty paragraph python-docx puts in a new cell."""
    paragraphs = cell.paragraphs
    if len(paragraphs) == 1 and not paragraphs[0].runs and not paragraphs[0].text:
        p = paragraphs[0]._element
        p.getparent().remove(p)


# --- public API ---------------------------------------------------------------


def export_docx(
    content_list: list[dict[str, Any]],
    images_dir: str | Path,
    output_path: str | Path,
    layout_mode: str = "proportional",
    page_sizes: list[Any] | None = None,
    middle_pages: list[Any] | None = None,
) -> Path:
    """Render a MinerU ``content_list`` into a .docx file.

    :param content_list: parsed ``*_content_list.json`` (a list of block dicts).
    :param images_dir: path to the ``images`` directory produced alongside the
        content list (used to resolve ``img_path`` references).
    :param output_path: destination path of the generated .docx file.
    :param layout_mode: ``"proportional"`` (default) preserves the relative
        geometry of the original page using each block's ``bbox``;
        ``"flow"`` reproduces the legacy full-width linear layout.
    :param page_sizes: optional list indexed by ``page_idx`` of ``[width,
        height]`` page sizes (e.g. from ``*_middle.json``); used to compute
        proportions and to size/orient the Word sheet. When absent or
        inconsistent with the bboxes, the reference box is estimated from the
        bbox extents.
    :param middle_pages: optional per-page layout info extracted from
        ``*_middle.json`` by :func:`_load_middle_pages` (page sizes plus text
        line heights); enables per-block font-size estimation. Supersedes
        ``page_sizes`` when both are given.
    :return: the resolved output path.
    """
    images_dir = Path(images_dir)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    proportional = str(layout_mode).lower() != "flow"

    blocks = [b for b in (content_list or []) if isinstance(b, dict)]

    if not proportional:
        ctx = _LayoutCtx(images_dir, proportional=False)
        section = document.sections[0]
        ctx.region_x0 = 0.0
        ctx.region_x1 = 1.0
        ctx.available_width = section.page_width - section.left_margin - section.right_margin
        for index, block in enumerate(content_list or []):
            _render_block(document, block, ctx, index)
        document.save(str(output_path))
        return output_path

    base_ctx = _LayoutCtx(images_dir, proportional=True)

    # Group blocks by page, preserving reading order; blocks without page_idx go
    # into their own trailing group.
    pages: dict[Any, list[dict[str, Any]]] = {}
    order: list[Any] = []
    for block in blocks:
        key = block.get("page_idx")
        if key not in pages:
            pages[key] = []
            order.append(key)
    for block in blocks:
        pages[block.get("page_idx")].append(block)

    def _middle_page_for(key: Any) -> dict[str, Any] | None:
        if middle_pages and isinstance(key, int) and 0 <= key < len(middle_pages):
            mp = middle_pages[key]
            if isinstance(mp, dict):
                return mp
        return None

    def _page_size_for(key: Any) -> tuple[float, float] | None:
        mp = _middle_page_for(key)
        candidate = mp.get("page_size") if mp else None
        if candidate is None and page_sizes and isinstance(key, int) and 0 <= key < len(page_sizes):
            candidate = page_sizes[key]
        if isinstance(candidate, (list, tuple)) and len(candidate) >= 2:
            try:
                return float(candidate[0]), float(candidate[1])
            except (TypeError, ValueError):
                return None
        return None

    ordered_keys = sorted([k for k in order if isinstance(k, int)]) + [k for k in order if not isinstance(k, int)]
    current_section = document.sections[0]
    for pos, key in enumerate(ordered_keys):
        page_blocks = pages[key]
        boxed = [b for b in page_blocks if _block_bbox(b) is not None]
        page_size = _page_size_for(key)
        region = _page_reference_box(boxed, page_size)

        # Word sheet = source page size (page_size when known, else the bbox
        # extents treated as points), clamped into the A5..A3 envelope with the
        # aspect ratio (and thus orientation) preserved. A new NEW_PAGE section
        # is only started when the sheet size actually changes; otherwise a
        # plain page break keeps the section count minimal.
        if page_size is not None:
            src_w, src_h = page_size
        else:
            src_w, src_h = region[2] - region[0], region[3] - region[1]
            if (
                src_w > 2
                and src_h > 2
                and max(src_w, src_h) / min(src_w, src_h) < _SQUARE_ASPECT_TOLERANCE
            ):
                src_w, src_h = _A4_SHORT_PT, _A4_LONG_PT
        page_w_pt, page_h_pt = _physical_page_pts(src_w, src_h)
        target_dims = (int(Pt(page_w_pt)), int(Pt(page_h_pt)))
        current_dims = (int(current_section.page_width), int(current_section.page_height))
        dims_close = (
            abs(target_dims[0] - current_dims[0]) <= current_dims[0] * _PAGE_DIMS_REUSE_TOL
            and abs(target_dims[1] - current_dims[1]) <= current_dims[1] * _PAGE_DIMS_REUSE_TOL
        )
        if pos == 0:
            _configure_page_section(current_section, page_w_pt, page_h_pt)
        elif dims_close:
            document.add_page_break()
        else:
            current_section = document.add_section(WD_SECTION.NEW_PAGE)
            _configure_page_section(current_section, page_w_pt, page_h_pt)

        usable_width = _section_usable_width(current_section)
        base_ctx.font_map, base_ctx.default_font_pt = _build_page_font_map(
            page_blocks, region, _middle_page_for(key), usable_width
        )

        if not boxed:
            # Nothing positionable: fall back to flow for this page's blocks.
            flow_ctx = base_ctx.clone_for(0.0, 1.0, usable_width)
            for b in page_blocks:
                _render_block(document, b, flow_ctx, 0)
        else:
            _render_region(document, boxed, region, usable_width, base_ctx, prefer_vertical=True, depth=0)
            # Render any bbox-less blocks on this page as trailing full-width flow.
            flow_ctx = base_ctx.clone_for(region[0], region[2], usable_width)
            for b in page_blocks:
                if _block_bbox(b) is None:
                    _render_block(document, b, flow_ctx, 0)

    document.save(str(output_path))
    return output_path


def find_content_list_json(result_dir: str | Path) -> Path | None:
    """Find the ``*_content_list.json`` file inside a MinerU parse result directory."""
    result_dir = Path(result_dir)
    if not result_dir.is_dir():
        return None
    matches = sorted(result_dir.glob(f"*{CONTENT_LIST_SUFFIX}"))
    return matches[0] if matches else None


def find_middle_json(result_dir: str | Path) -> Path | None:
    """Find the ``*_middle.json`` file inside a MinerU parse result directory."""
    result_dir = Path(result_dir)
    if not result_dir.is_dir():
        return None
    matches = sorted(result_dir.glob(f"*{MIDDLE_JSON_SUFFIX}"))
    return matches[0] if matches else None


def _load_middle_pages(middle_json_path: Path) -> list[Any] | None:
    """Extract per-page layout info from a MinerU ``*_middle.json`` file.

    Returns a list indexed by page of ``{"page_size": [w, h] | None,
    "blocks": [{"bbox": (x0, y0, x1, y1), "heights": [line heights...]}]}``,
    where line heights come from the text lines of each top-level block
    (``*_body`` leaves are excluded, see :func:`_collect_text_line_heights`).
    """
    try:
        with open(middle_json_path, "r", encoding="utf-8") as handle:
            middle = json.load(handle)
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning(f"Failed to read middle.json for page layout info: {exc}")
        return None
    page_info = middle.get("pdf_info") or middle.get("page_info")
    if not isinstance(page_info, list):
        return None
    result: list[Any] = []
    for page in page_info:
        if not isinstance(page, dict):
            result.append(None)
            continue
        page_size = page.get("page_size")
        if not isinstance(page_size, (list, tuple)):
            page_size = None
        mid_blocks: list[dict[str, Any]] = []
        raw_blocks = page.get("para_blocks") or page.get("preproc_blocks") or []
        if isinstance(raw_blocks, list):
            for raw in raw_blocks:
                if not isinstance(raw, dict):
                    continue
                bbox = _block_bbox(raw)
                heights = _collect_text_line_heights(raw)
                if bbox is not None:
                    mid_blocks.append({"bbox": bbox, "heights": heights})
        result.append({"page_size": page_size, "blocks": mid_blocks})
    return result


def export_docx_from_result_dir(
    result_dir: str | Path,
    output_path: str | Path | None = None,
    layout_mode: str = "proportional",
) -> Path:
    """Convenience wrapper: locate ``content_list.json``/``images`` in a parse result
    directory (e.g. ``<output_dir>/<name>/<parse_method>/``) and export a .docx file.

    If a ``*_middle.json`` file is present alongside the content list, its
    per-page ``page_size`` values are used to compute layout proportions and
    size/orient the Word sheets, and its text line heights drive the per-block
    font-size estimation.

    :param result_dir: directory containing ``*_content_list.json`` and an ``images``
        subdirectory (this is the per-document parse directory MinerU produces, as
        resolved by ``mineru.cli.output_paths.resolve_parse_dir``).
    :param output_path: destination .docx path; defaults to
        ``<result_dir>/<stem>.docx`` where ``<stem>`` is derived from the content
        list file name.
    :param layout_mode: ``"proportional"`` (default) or ``"flow"`` (see
        :func:`export_docx`).
    :return: the resolved output path.
    """
    result_dir = Path(result_dir)
    content_list_path = find_content_list_json(result_dir)
    if content_list_path is None:
        raise FileNotFoundError(f"No *{CONTENT_LIST_SUFFIX} file found under: {result_dir}")

    with open(content_list_path, "r", encoding="utf-8") as handle:
        content_list = json.load(handle)

    middle_pages = None
    middle_json_path = find_middle_json(result_dir)
    if middle_json_path is not None:
        middle_pages = _load_middle_pages(middle_json_path)

    images_dir = result_dir / "images"
    if output_path is None:
        stem = content_list_path.name[: -len(CONTENT_LIST_SUFFIX)]
        output_path = result_dir / f"{stem}.docx"

    return export_docx(
        content_list, images_dir, output_path, layout_mode=layout_mode, middle_pages=middle_pages
    )
